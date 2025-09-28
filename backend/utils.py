"""
工具函数
"""
import os
import uuid
import PyPDF2
import docx
from typing import Optional
from fastapi import UploadFile
from config import UPLOAD_DIR, MAX_FILE_SIZE, ALLOWED_EXTENSIONS

def generate_unique_filename(original_filename: str) -> str:
    """生成唯一文件名"""
    file_extension = os.path.splitext(original_filename)[1]
    unique_id = str(uuid.uuid4())
    return f"{unique_id}{file_extension}"

def validate_file(file: UploadFile) -> tuple[bool, str]:
    """验证上传文件"""
    # 检查文件大小
    if file.size and file.size > MAX_FILE_SIZE:
        return False, f"文件大小超过限制({MAX_FILE_SIZE / 1024 / 1024:.1f}MB)"
    
    # 检查文件扩展名
    file_extension = os.path.splitext(file.filename)[1].lower()
    if file_extension not in ALLOWED_EXTENSIONS:
        return False, f"不支持的文件格式，支持格式：{', '.join(ALLOWED_EXTENSIONS)}"
    
    return True, "文件验证通过"

def save_uploaded_file(file: UploadFile, filename: str) -> str:
    """保存上传的文件"""
    file_path = os.path.join(UPLOAD_DIR, filename)
    
    with open(file_path, "wb") as buffer:
        content = file.file.read()
        buffer.write(content)
    
    return file_path

def extract_text_from_pdf(file_path: str) -> str:
    """从PDF文件提取文本"""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text()
            
            return text.strip()
    except Exception as e:
        raise Exception(f"PDF文件解析失败: {str(e)}")

def extract_text_from_docx(file_path: str) -> str:
    """从DOCX文件提取文本"""
    try:
        doc = docx.Document(file_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text.strip()
    except Exception as e:
        raise Exception(f"DOCX文件解析失败: {str(e)}")

def extract_text_from_file(file_path: str) -> str:
    """根据文件类型提取文本内容"""
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    else:
        raise Exception(f"不支持的文件格式: {file_extension}")

def clean_text(text: str) -> str:
    """清理和格式化文本"""
    if not text:
        return ""
    
    # 移除多余的空白字符
    text = ' '.join(text.split())
    
    # 移除特殊字符
    import re
    text = re.sub(r'[^\w\s\u4e00-\u9fff.,;:!?()（）【】""''""''-]', '', text)
    
    return text.strip()

def calculate_match_score(skills_analysis: dict, experience_analysis: dict, education_analysis: dict) -> float:
    """计算综合匹配度评分"""
    # 技能匹配权重：40%
    skill_score = 0
    if skills_analysis.get("skill_scores"):
        skill_values = list(skills_analysis["skill_scores"].values())
        if skill_values:
            skill_score = sum(skill_values) / len(skill_values)
    
    # 经验匹配权重：40%
    experience_score = 0
    total_exp = experience_analysis.get("total_experience", 0)
    relevant_exp = experience_analysis.get("relevant_experience", 0)
    if total_exp > 0:
        experience_score = (relevant_exp / total_exp) * 100
    
    # 教育匹配权重：20%
    education_score = education_analysis.get("education_score", 0)
    
    # 综合评分
    final_score = (skill_score * 0.4) + (experience_score * 0.4) + (education_score * 0.2)
    
    return round(min(max(final_score, 0), 100), 1)

def format_duration(duration_str: str) -> float:
    """格式化工作时间为年数"""
    if not duration_str:
        return 0.0
    
    import re
    
    # 提取数字
    numbers = re.findall(r'\d+\.?\d*', duration_str)
    if not numbers:
        return 0.0
    
    duration = float(numbers[0])
    
    # 根据单位转换
    if '月' in duration_str or 'month' in duration_str.lower():
        return round(duration / 12, 1)
    elif '天' in duration_str or 'day' in duration_str.lower():
        return round(duration / 365, 1)
    else:
        return round(duration, 1)

def get_file_size_mb(file_path: str) -> float:
    """获取文件大小（MB）"""
    if not os.path.exists(file_path):
        return 0.0
    
    size_bytes = os.path.getsize(file_path)
    return round(size_bytes / 1024 / 1024, 2)
